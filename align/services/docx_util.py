import logging
from docx import Document
import re
from pathlib import Path
import ast


NIQQUD_RE = re.compile(r'[\u0591-\u05C7]')

def remove_nikud(text):
    text = NIQQUD_RE.sub('', text)
    
    # Step 3: Strip surrounding whitespace
    return text.strip()

def remove_parantheses(text):
    cleaned_text = re.sub(r'\([^)]*\)', '', text)
    return cleaned_text

def remove_marks_for_aligner(text):   
    
    clean_text = (text                 
                 .replace('-', ' ')
                 .replace('–', ' ')
                 .replace('—', ' ')
                 .replace(';', '.')
                 .replace(':', '.')
                 .replace('"', '')
                 .replace('״', '')
                 .replace('\'', '')
                 .replace('!', '.'))
                         
    return clean_text

def read_docx(file_path, nikud=False):
    if file_path == "":
        return ""
    doc = Document(file_path)
    full_text = []
    
    # Read paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # Remove starting " from the first word
        if text == '[' or text == '],':
            continue
        words = text.split()
        if words and words[0].startswith('"'):
            words[0] = words[0][1:]
        
        # Recombine words into a string
        text = ' '.join(words)   

        # Remove trailing " if it's at the end of the paragraph
        if text.endswith('",'):
            text = text[:-2]

        if text.endswith('"'):
            text = text[:-1]

        # Clean Hebrew text
        if not nikud:
            text = remove_nikud(text)            

        full_text.append(text)

    # Join with spaces to create continuous text
    return ' '.join(full_text)

def write_docx(output_file_path, text):
    try:
        document = Document()
        
        # Replace normal spaces after punctuation with non-breaking spaces
        text = text.replace('. ', '.\u00A0')
        text = text.replace(', ', ',\u00A0')
        text = text.replace('? ', '?\u00A0')
        text = text.replace('! ', '!\u00A0')
        text = text.replace(': ', ':\u00A0')
        text = text.replace('; ', ';\u00A0')
        
        document.add_paragraph(text)

        document.save(output_file_path)
    except Exception as e:
        print(f"An error occurred: {e}")

def read_file_content(filepath):
    with open(filepath, 'r', encoding='utf-8') as file:
        content = file.read()
    return content


def remove_special_chars_from_docx(file_path):
    # Read the document
    doc = Document(file_path)
    
    # Get all paragraphs text
    full_text = ' '.join(paragraph.text for paragraph in doc.paragraphs)
    
    # Clean the text:
    # 1. Remove quotes
    # 2. Remove square brackets
    # 3. Replace multiple spaces with single space
    # 4. Remove newlines
    clean_text = (full_text
                 .replace('"', '')
                 .replace('[', '')
                 .replace(']', '')
                 .replace('\n', ' '))
    
    # Use regex to replace 2 or more spaces with a single space
    clean_text = re.sub(r'\s{2,}', ' ', clean_text)
    
    # Optional: strip any leading/trailing whitespace
    clean_text = clean_text.strip()
    
    return clean_text

def save_cleaned_text_to_docx(cleaned_text, output_file_path):
    doc = Document()
    doc.add_paragraph(cleaned_text)
    doc.save(output_file_path)

def combine_start_text(text1, text2, last_words_count=40):
    try:       

        if last_words_count < 0:
            words = text2.split()
            # Convert negative count to positive and use it as starting index
            start_index = abs(last_words_count)
            # Return text2 starting from start_index (cutting off the beginning)
            return ' '.join(words[start_index:] if len(words) > start_index else words)        
        
        # Split into words and get last last_words_count words
        words = text1.split()
        last_words = ' '.join(words[(-1 * last_words_count):] if len(words) >= last_words_count else words)        

        # Combine content
        combined_content = f"{last_words} {text2}"
        
        return combined_content

    except Exception as e:
        print(f"Error processing files: {e}")
        return None
    

def combine_end_text(text1, text2, last_words_count=40):
    try:      
        

        if last_words_count < 0:           
            
            words = text1.split()
            # Convert negative count to positive and use it to remove words from the end
            remove_count = abs(last_words_count)
            # Return text1 with remove_count words removed from the end
            if len(words) > remove_count:
                return ' '.join(words[:-remove_count])
            else:
                return ''  # If we're removing more words than exist, return empty string       
        
        # Split text2 into words and get last last_words_count words from second doc
        words2 = text2.split()
        last_words_from_text2 = ' '.join(words2[:last_words_count] if len(words2) >= last_words_count else words2)

        # Combine content: full first doc + last words from second doc
        combined_content = f"{text1} {last_words_from_text2}"
        
        return combined_content

    except Exception as e:
        print(f"Error processing files: {e}")
        return None    
    

def read_statistics_from_srt(srt_file):
    result_probability_list = []

    with open(srt_file, 'r', encoding='utf-8-sig') as f:
        for line in f:
            line = line.strip()
            if line.startswith('[') and line.endswith(']'):
                try:
                    parsed = ast.literal_eval(line)
                    if isinstance(parsed, list) and all(isinstance(x, float) for x in parsed):
                        result_probability_list.append(parsed)
                except (SyntaxError, ValueError):
                    continue  # Skip lines that can't be parsed

    return result_probability_list